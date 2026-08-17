from pathlib import Path

from docx import Document as DocxDocument
from langchain_core.documents import Document


def load_docx_file(path: Path) -> list[Document]:
    document = DocxDocument(path)

    content = "\n".join(
        paragraph.text
        for paragraph in document.paragraphs
        if paragraph.text
    )

    return [
        Document(
            page_content=content,
            metadata={
                "source": str(path),
                "file_type": "docx",
            },
        )
    ]